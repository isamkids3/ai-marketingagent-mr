import os
import json
import logging
from dataclasses import dataclass
from typing import Any, Dict, List, Optional
from pathlib import Path
from dotenv import load_dotenv

# Load environment variables early from backend root directory
_dotenv_path = Path(__file__).resolve().parent.parent.parent / ".env"
load_dotenv(dotenv_path=_dotenv_path, override=True)

# pyrefly: ignore [missing-import]
from langchain_core.tools import tool
# pyrefly: ignore [missing-import]
from langchain_core.runnables import RunnableConfig

from app.agent.comfy_router import resolve_local_path

logger = logging.getLogger("agent_tools")
logger.setLevel(logging.INFO)

# Sandbox Environment Setup
BASE_WORKSPACE = Path(
    os.getenv("SHARED_WORKSPACE_ROOT", str(Path(__file__).parent.parent.parent.parent / "gen-content"))
).resolve()
try:
    os.makedirs(BASE_WORKSPACE, exist_ok=True)
except Exception as exc:
    logger.warning(f"Could not create BASE_WORKSPACE at {BASE_WORKSPACE}: {exc}")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def get_session_id_from_config(config: Optional[RunnableConfig]) -> str:
    """Extract session_id / thread_id from LangChain's RunnableConfig."""
    if not config:
        return "default"

    # Dict-style config
    if isinstance(config, dict):
        return config.get("configurable", {}).get("thread_id", "default")

    # Object-style config
    configurable = getattr(config, "configurable", {})
    if isinstance(configurable, dict):
        return configurable.get("thread_id", "default")

    return "default"


@dataclass
class SearchResult:
    """Structured result from a single web search hit."""
    title: str
    url: str
    snippet: str
    published_date: Optional[str]   # ISO 8601, e.g. "2026-06-15"
    source: str                      # e.g. "Reuters", "TechCrunch"
    relevance_score: float           # 0.0–1.0 from the search provider


# ---------------------------------------------------------------------------
# Search tool
# ---------------------------------------------------------------------------

@tool
def internet_search(
    query: str,
    max_results: int = 5,
    recency_days: Optional[int] = None,
) -> str:
    """Search for real-time marketing intelligence.

    Use this tool for competitor campaigns, industry trends, brand sentiment,
    campaign launches, share-of-voice benchmarks, earnings reports, consumer
    behaviour shifts, and any event that affects campaign context or timing. Add
    references to your response.

    Call this tool BEFORE forming a strategy when the request involves:
    - Competitor campaigns or positioning  ("what is Nike running right now")
    - Trend validation                     ("is plant-based still growing in SEA?")
    - Cultural moments or news hooks       ("World Cup 2026 final result")
    - Recent brand crises or PR events

    Call multiple times with focused queries to triangulate a topic — e.g.
    query the brand, then the category, then a key competitor separately.

    Args:
        query: Specific, objective search query (3–10 words).
               Good: "Nike Run Club campaign June 2026"
               Bad:  "tell me everything about Nike"
        max_results: Number of results to return (1–10).
                     Use 3 for a quick fact-check, 5–8 for trend research.
        recency_days: Restrict results to the last N days.
                      Use 7 for breaking news, 90 for trend analysis,
                      None for no restriction.

    Returns:
        JSON array of SearchResult objects sorted by relevance_score desc.
        Each entry contains: title, url, snippet, published_date, source,
        relevance_score. Parse the JSON and reason over individual results —
        do not treat the whole payload as a single fact.
    """
    # pyrefly: ignore [missing-import]
    from tavily import TavilyClient
    client = TavilyClient(api_key=os.environ["TAVILY_API_KEY"])
    raw = client.search(
        query,
        max_results=max_results,
        days=recency_days,
        include_raw_content=False,
    )
    results = [
        SearchResult(
            title=r["title"],
            url=r["url"],
            snippet=r["content"],
            published_date=r.get("published_date"),
            source=r.get("source", "Unknown"),
            relevance_score=r.get("score", 1.0),
        )
        for r in raw["results"]
    ]
    return json.dumps([vars(r) for r in results], ensure_ascii=False)


# ---------------------------------------------------------------------------
# Document reading tool
# ---------------------------------------------------------------------------

@tool
def read_user_document_tool(file_path: str) -> str:
    """Extract plain-text content from a user-uploaded document.

    Use this tool when you need to read briefs, scripts, long prompts, or
    structured descriptions from an uploaded file in the workspace.

    Supported formats: .txt, .pdf, .docx

    Args:
        file_path: Sandbox path to the document
                   (e.g. /workspace/brief.pdf).
    """
    resolved_path = resolve_local_path(file_path)
    logger.info(f"read_user_document_tool: reading {resolved_path!r}")

    if not os.path.exists(resolved_path):
        return (
            f"Error: file not found at '{file_path}' "
            f"(resolved: '{resolved_path}'). Please verify the path."
        )

    ext = os.path.splitext(resolved_path)[1].lower()

    try:
        if ext == ".txt":
            with open(resolved_path, "r", encoding="utf-8", errors="ignore") as fh:
                return fh.read().strip()

        elif ext == ".pdf":
            # pyrefly: ignore [missing-import]
            import pypdf
            reader = pypdf.PdfReader(resolved_path)
            pages = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                pages.append(f"--- Page {i + 1} ---\n{text.strip()}")
            return "\n\n".join(pages)

        elif ext == ".docx":
            # pyrefly: ignore [missing-import]
            import docx
            doc = docx.Document(resolved_path)
            paragraphs = [p.text for p in doc.paragraphs]
            table_rows = []
            for table in doc.tables:
                for row in table.rows:
                    table_rows.append(
                        " | ".join(cell.text.strip() for cell in row.cells)
                    )
            content = "\n".join(paragraphs)
            if table_rows:
                content += "\n\n--- Document Tables ---\n" + "\n".join(table_rows)
            return content.strip()

        else:
            return (
                f"Error: unsupported format '{ext}'. "
                "Only .txt, .pdf, and .docx files are supported."
            )

    except Exception as exc:
        logger.error(
            f"Error reading {resolved_path!r}: {exc}", exc_info=True
        )
        return f"Error: could not read the document. Details: {exc}"


# ---------------------------------------------------------------------------
# Sandbox Write Tool
# ---------------------------------------------------------------------------

@tool
def write_file_to_sandbox(
    filename: str,
    content: str,
    append: bool = False,
    config: RunnableConfig = None,
) -> str:
    """Write or append text content to the secure sandbox.

    Use this tool when you need to save campaign plans, markdown files,
    code snippets, or text generated during your tasks.

    For very long files that exceed output limits, you can write the first
    chunk with `append=False` (default), and then append subsequent chunks
    with `append=True`.

    The `thread_id` is automatically populated from the system context; you do
    not need to guess or provide it. After saving, the tool will return a 
    web-accessible URL pointer to the file. You can retrieve and use this URL
    in your responses to the user.

    Args:
        filename: The name of the file to create or append to (e.g. "campaign_plan.md").
        content: The text content to write or append into the file.
        append: If True, appends content to the file if it exists. If False, overwrites or creates the file.
    """
    from app.agent.orchestrator import active_session_id
    thread_id = active_session_id.get()
    
    if not thread_id or thread_id == "default":
        thread_id = get_session_id_from_config(config)
        
    if not thread_id:
        thread_id = "anonymous"

    thread_dir = (BASE_WORKSPACE / thread_id).resolve()
    os.makedirs(thread_dir, exist_ok=True)

    file_path = (thread_dir / filename).resolve()

    # Path Traversal protection
    if not str(file_path).startswith(str(thread_dir)):
        raise ValueError("Path traversal detected: invalid filename.")

    try:
        mode = "a" if append else "w"
        with open(file_path, mode, encoding="utf-8") as f:
            f.write(content)
    except Exception as exc:
        logger.error(f"Error writing to sandbox {file_path!r}: {exc}", exc_info=True)
        return f"Error: could not write file. Details: {exc}"

    action = "appended to" if append else "saved"
    return f"File {action} successfully. Access it at: /sandbox/{thread_id}/{filename}"


# ---------------------------------------------------------------------------
# Sandbox PDF Generator Tool
# ---------------------------------------------------------------------------

@tool
def generate_pdf_in_sandbox(
    filename: str,
    content: str,
    title: Optional[str] = None,
    config: RunnableConfig = None,
) -> str:
    """Generate a professionally-styled PDF document and save it in the sandbox.

    Use this tool when you need to output a formatted report, brief, campaign
    plan, copy sheet, or proposal as a PDF file for the user.

    The content parameter accepts Markdown formatting, which will be compiled
    into structured PDF sections:
      - Headings: `# Heading 1`, `## Heading 2`, `### Heading 3`
      - Paragraphs: Standard text flow
      - Bullet points: `- Step 1` or `* Step 2`
      - Numbered lists: `1. First item`
      - Bold/Italic formatting: `**bold**` or `*italic*`
      - Image embedding: `![Alt Text](/sandbox/{session_id}/images/generated_{id}.png)` or HTTP web image URLs. **Note**: Always generate your images first to get their `/sandbox/` paths before calling this tool to compile the PDF.

    The `thread_id` is automatically populated from the system context. After
    saving, the tool will return a web-accessible URL pointer to the PDF.

    Args:
        filename: The name of the file to save (must end in '.pdf', e.g. "campaign_brief.pdf").
        content: The Markdown-formatted text to render into the PDF.
        title: Optional title to place at the top of the first page and in the page headers.
    """
    from app.agent.orchestrator import active_session_id
    thread_id = active_session_id.get()
    
    if not thread_id or thread_id == "default":
        thread_id = get_session_id_from_config(config)
        
    if not thread_id:
        thread_id = "anonymous"

    # Enforce .pdf extension
    if not filename.lower().endswith(".pdf"):
        filename = f"{filename}.pdf"

    thread_dir = (BASE_WORKSPACE / thread_id).resolve()
    os.makedirs(thread_dir, exist_ok=True)

    file_path = (thread_dir / filename).resolve()

    # Path Traversal protection
    if not str(file_path).startswith(str(thread_dir)):
        raise ValueError("Path traversal detected: invalid filename.")

    try:
        import re
        import html
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib.colors import HexColor

        temp_files_to_clean = []

        # Helper to resolve sandbox / local / HTTP image paths
        def resolve_image_path(path_str: str) -> Optional[str]:
            import httpx
            path_str = path_str.strip()
            
            # If it starts with /sandbox/
            if path_str.startswith("/sandbox/"):
                relative_path = path_str.replace("/sandbox/", "", 1)
                full_path = (BASE_WORKSPACE / relative_path).resolve()
                if full_path.exists():
                    return str(full_path)
            # If it is a remote HTTP URL
            elif path_str.startswith("http://") or path_str.startswith("https://"):
                try:
                    with httpx.Client(timeout=15.0) as client:
                        resp = client.get(path_str)
                        if resp.status_code == 200:
                            import tempfile
                            temp_dir = Path(tempfile.gettempdir())
                            temp_file = temp_dir / f"downloaded_{abs(hash(path_str))}.png"
                            with open(temp_file, "wb") as f:
                                f.write(resp.content)
                            temp_files_to_clean.append(temp_file)
                            return str(temp_file)
                except Exception as e:
                    logger.error(f"Failed to download image {path_str}: {e}")
            # If it is a relative path in the thread directory
            elif not path_str.startswith("/") and not path_str.startswith("http"):
                full_path = (thread_dir / path_str).resolve()
                if full_path.exists():
                    return str(full_path)
            # If it is already a direct absolute path
            elif os.path.exists(path_str):
                return path_str
            return None

        # Document setup with 0.75in left/right margins, and larger top margin for header line
        doc = SimpleDocTemplate(
            str(file_path),
            pagesize=letter,
            leftMargin=54,
            rightMargin=54,
            topMargin=72,
            bottomMargin=54,
        )

        # Style sheet setup
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            'DocTitle',
            parent=styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=22,
            leading=26,
            textColor=HexColor('#1A365D'),
            spaceAfter=15,
            alignment=TA_LEFT
        )
        
        h1_style = ParagraphStyle(
            'DocH1',
            parent=styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=15,
            leading=18,
            textColor=HexColor('#1A365D'),
            spaceBefore=14,
            spaceAfter=6,
            keepWithNext=True
        )
        
        h2_style = ParagraphStyle(
            'DocH2',
            parent=styles['Heading3'],
            fontName='Helvetica-Bold',
            fontSize=12,
            leading=15,
            textColor=HexColor('#2B6CB0'),
            spaceBefore=12,
            spaceAfter=4,
            keepWithNext=True
        )
        
        h3_style = ParagraphStyle(
            'DocH3',
            parent=styles['Heading4'],
            fontName='Helvetica-Bold',
            fontSize=10.5,
            leading=13,
            textColor=HexColor('#4A5568'),
            spaceBefore=8,
            spaceAfter=3,
            keepWithNext=True
        )

        body_style = ParagraphStyle(
            'DocBody',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            leading=14.5,
            textColor=HexColor('#2D3748'),
            spaceAfter=8
        )

        bullet_style = ParagraphStyle(
            'DocBullet',
            parent=body_style,
            leftIndent=20,
            bulletIndent=8,
            spaceAfter=4
        )

        # Markdown inline parsing helper
        def md_to_reportlab_html(text: str) -> str:
            # Escape XML entities safely
            text = html.escape(text)
            # Convert bold **text** or __text__ to <b>text</b>
            text = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", text)
            text = re.sub(r"__(.*?)__", r"<b>\1</b>", text)
            # Convert italic *text* or _text_ to <i>text</i>
            text = re.sub(r"\*(.*?)\*", r"<i>\1</i>", text)
            text = re.sub(r"_(.*?)_", r"<i>\1</i>", text)
            # Convert inline code `code` to Courier font
            text = re.sub(r"`(.*?)`", r'<font name="Courier">\1</font>', text)
            return text

        # Parse blocks
        raw_blocks = content.split("\n\n")
        story = []

        # Document Header block (Title on the first page)
        doc_title = title or filename.rsplit('.', 1)[0].replace('_', ' ').title()
        story.append(Paragraph(html.escape(doc_title), title_style))
        story.append(Spacer(1, 10))

        for raw_block in raw_blocks:
            block = raw_block.strip()
            if not block:
                continue
                
            # Check for Headings
            if block.startswith("# "):
                text = block[2:].strip()
                story.append(Paragraph(md_to_reportlab_html(text), h1_style))
            elif block.startswith("## "):
                text = block[3:].strip()
                story.append(Paragraph(md_to_reportlab_html(text), h2_style))
            elif block.startswith("### "):
                text = block[4:].strip()
                story.append(Paragraph(md_to_reportlab_html(text), h3_style))
            elif block.startswith("#### "):
                text = block[5:].strip()
                story.append(Paragraph(md_to_reportlab_html(text), h3_style))
            # Check if block is a list
            elif block.startswith(("- ", "* ", "• ")) or re.match(r"^\d+\.\s", block):
                lines = block.split("\n")
                for line in lines:
                    line_str = line.strip()
                    if not line_str:
                        continue
                    if line_str.startswith(("- ", "* ", "• ")):
                        text = line_str[2:].strip()
                        story.append(Paragraph(md_to_reportlab_html(text), bullet_style, bulletText='•'))
                    else:
                        match = re.match(r"^(\d+\.)\s(.*)", line_str)
                        if match:
                            bullet_num = match.group(1)
                            text = match.group(2).strip()
                            story.append(Paragraph(md_to_reportlab_html(text), bullet_style, bulletText=bullet_num))
                        else:
                            story.append(Paragraph(md_to_reportlab_html(line_str), bullet_style, bulletText=''))
            else:
                # Standard paragraph or line-by-line check for images/text
                lines = block.split("\n")
                current_text_lines = []
                
                for line in lines:
                    line_str = line.strip()
                    if line_str.startswith("![") and line_str.endswith(")"):
                        # If we have accumulated text before this image, render it first
                        if current_text_lines:
                            cleaned_text = " ".join(current_text_lines)
                            story.append(Paragraph(md_to_reportlab_html(cleaned_text), body_style))
                            current_text_lines = []
                            
                        # Process image
                        match = re.match(r"^!\[(.*?)\]\((.*?)\)$", line_str)
                        if match:
                            alt_text = match.group(1)
                            img_path_str = match.group(2)
                            resolved_path = resolve_image_path(img_path_str)
                            if resolved_path:
                                try:
                                    from reportlab.platypus import Image as RLImage
                                    from PIL import Image as PILImage
                                    
                                    with PILImage.open(resolved_path) as pil_img:
                                        w, h = pil_img.size
                                    
                                    # Scale maintaining aspect ratio (fit within margins)
                                    max_w = 504
                                    max_h = 350
                                    scale = min(max_w / w, max_h / h, 1.0)
                                    new_w = w * scale
                                    new_h = h * scale
                                    
                                    rl_img = RLImage(resolved_path, width=new_w, height=new_h)
                                    rl_img.hAlign = 'CENTER'
                                    story.append(rl_img)
                                    story.append(Spacer(1, 8))
                                except Exception as e:
                                    logger.error(f"Failed to embed image {resolved_path}: {e}")
                                    story.append(Paragraph(f"<i>(Failed to load image: {alt_text})</i>", body_style))
                            else:
                                story.append(Paragraph(f"<i>(Image not found: {alt_text})</i>", body_style))
                    else:
                        current_text_lines.append(line_str)
                
                # Render any remaining text
                if current_text_lines:
                    cleaned_text = " ".join(current_text_lines)
                    story.append(Paragraph(md_to_reportlab_html(cleaned_text), body_style))

        # Canvas drawing callbacks for running header/footer
        def draw_later_page(canvas, document):
            canvas.saveState()
            # Draw header title
            canvas.setFont('Helvetica-Bold', 8.5)
            canvas.setFillColor(HexColor('#4A5568'))
            canvas.drawString(54, 750, doc_title.upper())
            
            # Draw header line
            canvas.setStrokeColor(HexColor('#E2E8F0'))
            canvas.setLineWidth(0.5)
            canvas.line(54, 742, letter[0] - 54, 742)
            
            # Draw footer page number
            canvas.setFont('Helvetica', 8.5)
            canvas.setFillColor(HexColor('#718096'))
            canvas.drawRightString(letter[0] - 54, 36, f"Page {document.page}")
            canvas.restoreState()

        def draw_first_page(canvas, document):
            canvas.saveState()
            # Draw footer page number only on the first page
            canvas.setFont('Helvetica', 8.5)
            canvas.setFillColor(HexColor('#718096'))
            canvas.drawRightString(letter[0] - 54, 36, f"Page {document.page}")
            canvas.restoreState()

        # Build document
        doc.build(story, onFirstPage=draw_first_page, onLaterPages=draw_later_page)

    except Exception as exc:
        logger.error(f"Error generating PDF in sandbox {file_path!r}: {exc}", exc_info=True)
        return f"Error: could not generate PDF. Details: {exc}"
    finally:
        # Clean up temporary downloaded files
        for temp_file in temp_files_to_clean:
            try:
                if temp_file.exists():
                    os.remove(temp_file)
            except Exception as e:
                logger.error(f"Failed to clean up temp file {temp_file}: {e}")

    return f"PDF generated successfully. Access it at: /sandbox/{thread_id}/{filename}"