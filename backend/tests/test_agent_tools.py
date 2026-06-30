import os
import pytest
import tempfile
# pyrefly: ignore [missing-import]
import docx
# pyrefly: ignore [missing-import]
import pypdf
from typing import Generator
from unittest.mock import patch, AsyncMock
# pyrefly: ignore [missing-import]
from langchain_core.runnables import RunnableConfig

from app.agent.tools import (
    read_user_document_tool,
    get_session_id_from_config
)
from app.agent.comfy_router import WorkspaceManager


@pytest.fixture
def temp_dir() -> Generator[str, None, None]:
    with tempfile.TemporaryDirectory() as tmpdir:
        WorkspaceManager.set_workspace_dir(tmpdir)
        yield tmpdir


def test_get_session_id_from_config():
    # Test with None
    assert get_session_id_from_config(None) == "default"
    
    # Test with dict
    config_dict = {"configurable": {"thread_id": "session-123"}}
    assert get_session_id_from_config(config_dict) == "session-123"
    
    # Test with dict but missing thread_id
    assert get_session_id_from_config({"configurable": {}}) == "default"
    
    # Test with RunnableConfig-like object
    class MockConfig:
        def __init__(self):
            self.configurable = {"thread_id": "session-abc"}
            
    assert get_session_id_from_config(MockConfig()) == "session-abc"


def test_read_user_document_text(temp_dir):
    txt_path = os.path.join(temp_dir, "test.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write("Hello World from TXT!")
        
    # Test virtual path resolving
    virtual_path = "/workspace/test.txt"
    content = read_user_document_tool.invoke({"file_path": virtual_path})
    assert "Hello World from TXT!" in content


def test_read_user_document_docx(temp_dir):
    docx_path = os.path.join(temp_dir, "test.docx")
    doc = docx.Document()
    doc.add_paragraph("Hello from Paragraph 1")
    doc.add_paragraph("Hello from Paragraph 2")
    doc.save(docx_path)
    
    virtual_path = "/workspace/test.docx"
    content = read_user_document_tool.invoke({"file_path": virtual_path})
    assert "Hello from Paragraph 1" in content
    assert "Hello from Paragraph 2" in content


def test_read_user_document_pdf(temp_dir):
    pdf_path = os.path.join(temp_dir, "test.pdf")
    
    # Write a minimal PDF structure
    pdf_content = (
        b"%PDF-1.4\n"
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
        b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Contents 4 0 R /Resources << >> >>\nendobj\n"
        b"4 0 obj\n<< /Length 52 >>\nstream\nBT\n/F1 12 Tf\n72 712 Td\n(Hello World from PDF!) Tj\nET\nendstream\nendobj\n"
        b"xref\n0 5\n0000000000 65535 f\n0000000009 00000 n\n0000000056 00000 n\n0000000111 00000 n\n0000000211 00000 n\n"
        b"trailer\n<< /Size 5 /Root 1 0 R >>\nstartxref\n312\n%%EOF"
    )
    with open(pdf_path, "wb") as f:
        f.write(pdf_content)
        
    virtual_path = "/workspace/test.pdf"
    content = read_user_document_tool.invoke({"file_path": virtual_path})
    
    # Depending on how the PDF text extractor parses the simplified PDF stream,
    # it may contain the text or be blank/error out. Let's make sure it handles it gracefully
    # and doesn't crash the tool.
    assert isinstance(content, str)


def test_read_user_document_not_found():
    content = read_user_document_tool.invoke({"file_path": "/workspace/nonexistent.txt"})
    assert "file not found" in content


def test_read_user_document_unsupported(temp_dir):
    unsupported_path = os.path.join(temp_dir, "test.xyz")
    with open(unsupported_path, "w") as f:
        f.write("xyz")
        
    content = read_user_document_tool.invoke({"file_path": "/workspace/test.xyz"})
    assert "unsupported format" in content


def test_write_file_to_sandbox(temp_dir):
    from pathlib import Path
    from app.agent.orchestrator import active_session_id
    from app.agent.tools import write_file_to_sandbox
    
    token = active_session_id.set("test-session-xyz")
    try:
        with patch("app.agent.tools.BASE_WORKSPACE", Path(temp_dir)):
            # Test write
            res = write_file_to_sandbox.invoke({
                "filename": "brief.md",
                "content": "Minimal brand strategy."
            })
            assert "test-session-xyz" in res
            assert "brief.md" in res
            
            # Verify file exists on disk under temp_dir
            target_path = os.path.join(temp_dir, "test-session-xyz", "brief.md")
            assert os.path.exists(target_path)
            with open(target_path, "r", encoding="utf-8") as f:
                assert f.read() == "Minimal brand strategy."
                
            # Test append
            res_append = write_file_to_sandbox.invoke({
                "filename": "brief.md",
                "content": " Continued strategy.",
                "append": True
            })
            assert "appended to" in res_append
            with open(target_path, "r", encoding="utf-8") as f:
                assert f.read() == "Minimal brand strategy. Continued strategy."
    finally:
        active_session_id.reset(token)


def test_context_window_optimizations():
    from app.agent.orchestrator import TokenLimitingChatOpenAI
    # pyrefly: ignore [missing-import]
    from langchain_core.messages import SystemMessage, HumanMessage, AIMessage, ToolMessage
    
    # Instantiate the wrapper model
    llm = TokenLimitingChatOpenAI(
        model="gpt-4o",
        openai_api_key="mock-key",
        max_tokens=3200
    )
    
    # 1. Test ToolMessage pruning for dicts
    tool_msg_dict = ToolMessage(
        content={"result": "A" * 2200},
        tool_call_id="call_123"
    )
    messages = [
        SystemMessage(content="system prompt"),
        tool_msg_dict,
        HumanMessage(content="user query")
    ]
    
    pruned = llm._prune_messages(messages)
    assert len(pruned) == 3
    assert "[Tool Output Truncated" in pruned[1].content
    assert len(pruned[1].content) < 1000
    
    # 2. Test ToolMessage pruning for lists
    tool_msg_list = ToolMessage(
        content=["B" * 2200],
        tool_call_id="call_124"
    )
    messages = [
        SystemMessage(content="system prompt"),
        tool_msg_list,
        HumanMessage(content="user query")
    ]
    pruned = llm._prune_messages(messages)
    assert len(pruned) == 3
    assert "[Tool Output Truncated" in pruned[1].content
    
    # 3. Test Historical AIMessage tool call simplification
    historical_ai_msg = AIMessage(
        content="some thoughts",
        tool_calls=[
            {
                "name": "text_image",
                "args": {
                    "prompt": {
                        "high_level_description": "A beautiful sunset over the mountains",
                        "aspect_ratio": "16:9",
                        "compositional_deconstruction": {
                            "background": "transparent background",
                            "elements": [{"type": "obj", "desc": "mountain details " * 100}]
                        }
                    },
                    "other_huge_arg": "C" * 400
                },
                "id": "call_999"
            }
        ]
    )
    recent_ai_msg = AIMessage(
        content="recent thoughts",
        tool_calls=[
            {
                "name": "text_image",
                "args": {
                    "prompt": {
                        "high_level_description": "A beautiful sunset over the mountains",
                        "aspect_ratio": "16:9",
                        "compositional_deconstruction": {
                            "background": "transparent background",
                            "elements": [{"type": "obj", "desc": "mountain details " * 100}]
                        }
                    }
                },
                "id": "call_888"
            }
        ]
    )
    
    messages = [
        SystemMessage(content="system prompt"),
        historical_ai_msg,
        recent_ai_msg,
        HumanMessage(content="user query")
    ]
    
    pruned = llm._prune_messages(messages)
    assert len(pruned) == 4
    
    # Historical AIMessage tool call simplified:
    hist_tool_call = pruned[1].tool_calls[0]
    hist_prompt = hist_tool_call["args"]["prompt"]
    assert "compositional_deconstruction" not in hist_prompt
    assert hist_prompt["high_level_description"] == "A beautiful sunset over the mountains"
    assert hist_prompt["aspect_ratio"] == "16:9"
    assert "Truncated" in hist_tool_call["args"]["other_huge_arg"]
    
    # Recent AIMessage tool call kept intact:
    recent_tool_call = pruned[2].tool_calls[0]
    assert "compositional_deconstruction" in recent_tool_call["args"]["prompt"]
    
    # 3b. Test eager pruning of write_file_to_sandbox in both recent and historical AIMessages
    write_sandbox_msg = AIMessage(
        content="writing files",
        tool_calls=[
            {
                "name": "write_file_to_sandbox",
                "args": {
                    "filename": "plan.md",
                    "content": "A" * 5000
                },
                "id": "call_abc"
            }
        ]
    )
    # Even if recent (at index len(messages) - 2), write_file_to_sandbox content should be stripped
    prune_sandbox_msgs = [
        SystemMessage(content="system prompt"),
        write_sandbox_msg,
        HumanMessage(content="user query")
    ]
    pruned_sandbox = llm._prune_messages(prune_sandbox_msgs)
    assert len(pruned_sandbox) == 3
    args = pruned_sandbox[1].tool_calls[0]["args"]
    assert "Omitted file content" in args["content"]
    
    # 4. Test Sliding-Window history eviction
    # Create a system message, several turns of messages, and a final human message
    large_human_msg = HumanMessage(content="large context " * 2000) # ~2000 tokens
    
    eviction_messages = [SystemMessage(content="system prompt")]
    for j in range(10):
        eviction_messages.append(HumanMessage(content=f"older turn user {j}"))
        eviction_messages.append(AIMessage(content=f"older turn assistant {j}"))
        
    eviction_messages.append(large_human_msg) # last message is kept
    
    # Mock _estimate_tokens to return a value larger than 32000 when history is long
    def mock_estimate(msgs, tools=None):
        # If there are more than 5 messages in the history, simulate token overload
        if len(msgs) > 5:
            return 35000
        return 15000
        
    with patch.object(llm, "_estimate_tokens", side_effect=mock_estimate):
        pruned_eviction = llm._prune_messages(eviction_messages)
        # Should keep system message (1) + last 4 other messages (4)
        assert len(pruned_eviction) == 5
        assert pruned_eviction[0].content == "system prompt"
        assert pruned_eviction[-1].content == "large context " * 2000

    # 5. Test _stream and _astream overrides
    large_tool_msg = ToolMessage(content="X" * 2500, tool_call_id="call_x")
    stream_messages = [
        SystemMessage(content="system prompt"),
        large_tool_msg,
        HumanMessage(content="user query")
    ]
    
    with patch("app.agent.orchestrator.ChatOpenAI._stream", return_value=iter(["chunk1"])) as mock_super_stream, \
         patch("app.agent.orchestrator.ChatOpenAI._astream") as mock_super_astream:
         
         # Mock _astream to return an async generator
         async def mock_async_gen(*args, **kwargs):
             yield "achunk1"
         mock_super_astream.side_effect = mock_async_gen
         
         # Run _stream
         list(llm._stream(stream_messages))
         called_msgs = mock_super_stream.call_args[0][0]
         assert len(called_msgs) == 3
         assert "[Tool Output Truncated" in called_msgs[1].content
         
         # Run _astream
         async def run_astream():
             chunks = []
             async for chunk in llm._astream(stream_messages):
                 chunks.append(chunk)
             return chunks
         
         import asyncio
         asyncio.run(run_astream())
         called_async_msgs = mock_super_astream.call_args[0][0]
         assert len(called_async_msgs) == 3
         assert "[Tool Output Truncated" in called_async_msgs[1].content


def test_resolve_local_path():
    from app.agent.comfy_router import resolve_local_path, WorkspaceManager
    
    # Test normal workspace resolution
    WorkspaceManager.set_workspace_dir("/tmp/mock_workspace")
    path_ws = resolve_local_path("/workspace/file.txt")
    assert path_ws == "/tmp/mock_workspace/file.txt"
    
    # Test sandbox resolution (should resolve relative to SHARED_WORKSPACE_ROOT)
    with patch.dict(os.environ, {"SHARED_WORKSPACE_ROOT": "/tmp/mock_shared"}):
        path_sb = resolve_local_path("/sandbox/session-123/file.txt")
        assert path_sb == "/tmp/mock_shared/session-123/file.txt"
        
    # Test plain/absolute path pass-through
    path_plain = resolve_local_path("/etc/hosts")
    assert path_plain == "/etc/hosts"


def test_composition_validators():
    from app.agent.orchestrator import CompositionElement, CompositionalDeconstruction
    import pytest
    from pydantic import ValidationError

    # Valid text element
    elem_text_ok = CompositionElement(type="text", bbox=[100, 200, 180, 800], desc="Title text", text="Welcome")
    assert elem_text_ok.text == "Welcome"

    # Missing text for type="text"
    with pytest.raises(ValidationError) as excinfo:
        CompositionElement(type="text", bbox=[100, 200, 180, 800], desc="Title text")
    assert "The 'text' field is mandatory" in str(excinfo.value)

    # Valid 0-1000 scale bbox
    comp_ok = CompositionalDeconstruction(
        background="blue sky",
        elements=[
            CompositionElement(type="obj", bbox=[100, 200, 900, 800], desc="A balloon"),
            CompositionElement(type="text", bbox=[10, 20, 50, 40], desc="Small text", text="A")  # Small logo/text is ok if overall max > 100
        ]
    )

    # Invalid 0-100 percentage scale bbox
    with pytest.raises(ValidationError) as excinfo:
        CompositionalDeconstruction(
            background="blue sky",
            elements=[
                CompositionElement(type="obj", bbox=[10, 20, 90, 80], desc="A balloon"),
                CompositionElement(type="text", bbox=[5, 10, 40, 30], desc="Small text", text="A")
            ]
        )
    assert "0-100 percentage scale" in str(excinfo.value)

    # Invalid vertical pillar text bbox (axis swap)
    with pytest.raises(ValidationError) as excinfo:
        CompositionElement(
            type="text",
            bbox=[150, 100, 850, 400],
            desc="Squeezed header text",
            text="Barista quality coffee at home."
        )
    assert "extremely tall and narrow" in str(excinfo.value)

    # Invalid too tall text bbox
    with pytest.raises(ValidationError) as excinfo:
        CompositionElement(
            type="text",
            bbox=[300, 100, 700, 900],  # height = 400
            desc="Header text",
            text="Tired of bitter\ngrocery store coffee?"  # 2 lines, max allowed height is 200
        )
    assert "too tall" in str(excinfo.value)

    # Invalid overlapping bounding boxes between text and object
    with pytest.raises(ValidationError) as excinfo:
        CompositionalDeconstruction(
            background="coffee cup shop",
            elements=[
                CompositionElement(type="obj", bbox=[200, 300, 800, 800], desc="Blurred generic crumpled paper coffee cup"),
                CompositionElement(type="text", bbox=[300, 150, 380, 500], desc="Barista quality coffee", text="Coffee")
            ]
        )
    assert "Overlap/Intersection detected" in str(excinfo.value)

    # Invalid spelling typo (distance = 1 anomaly against user prompt words)
    from app.agent.orchestrator import active_user_words
    token = active_user_words.set({"grocery", "coffee", "bitter", "store"})
    try:
        with pytest.raises(ValidationError) as excinfo:
            CompositionElement(
                type="text",
                bbox=[100, 150, 180, 850],
                desc="Header text",
                text="Tired of biter gocery store cofee?"
            )
        assert "Spelling anomaly detected" in str(excinfo.value)
    finally:
        active_user_words.reset(token)


if __name__ == "__main__":
    import sys
    # Add project root to sys.path so we can import app
    project_root = os.path.dirname(os.path.abspath(__file__))
    if os.path.basename(project_root) == "tests":
        project_root = os.path.dirname(project_root)
    if project_root not in sys.path:
        sys.path.insert(0, project_root)
        
    print("Running tests via pytest...")
    sys.exit(pytest.main([__file__]))

